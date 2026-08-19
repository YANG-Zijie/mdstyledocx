from __future__ import annotations

import os
import re
import unicodedata
from collections.abc import Mapping
from dataclasses import dataclass, field, replace
from datetime import UTC, datetime
from io import BytesIO
from typing import Any

from docx import Document as WordDocument
from docx.document import Document as WordprocessingDocument
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Twips
from lxml import etree

from mdstyledocx.model import Block, Document, ImageSpan, InlineElement, InlineSpan
from mdstyledocx.presets import Preset, Style

TEMPLATE_TOKEN_RE = re.compile(r"\{(page|pages|title|date)\}")
PAGE_FIELD_CODES = {"page": "PAGE", "pages": "NUMPAGES"}
VML_NS = "urn:schemas-microsoft-com:vml"
OFFICE_NS = "urn:schemas-microsoft-com:office:office"
WORD_2003_NS = "urn:schemas-microsoft-com:office:word"


@dataclass
class BuildState:
    preset: Preset
    heading_counters: dict[int, int] = field(default_factory=dict)


def build_docx(document: Document, preset: Preset) -> bytes:
    word_document = WordDocument()
    _configure_document(word_document, preset)
    _configure_page_content(word_document, document, preset)
    _set_core_properties(word_document, _document_title(document))

    state = BuildState(preset=preset)
    for index, block in enumerate(document.blocks):
        next_block = (
            document.blocks[index + 1]
            if index + 1 < len(document.blocks)
            else None
        )
        _append_block(word_document, block, state, next_block=next_block)

    buffer = BytesIO()
    word_document.save(buffer)
    return buffer.getvalue()


def _configure_document(word_document: WordprocessingDocument, preset: Preset) -> None:
    section = word_document.sections[0]
    section.page_width = Twips(preset.page.width)
    section.page_height = Twips(preset.page.height)
    section.top_margin = Twips(preset.page.margin_top)
    section.right_margin = Twips(preset.page.margin_right)
    section.bottom_margin = Twips(preset.page.margin_bottom)
    section.left_margin = Twips(preset.page.margin_left)
    section.header_distance = Twips(preset.page.header)
    section.footer_distance = Twips(preset.page.footer)
    section.gutter = Twips(preset.page.gutter)


def _configure_page_content(
    word_document: WordprocessingDocument, document: Document, preset: Preset
) -> None:
    header_content = _page_region_content(document.metadata.get("header"), "header")
    footer_content = _page_region_content(document.metadata.get("footer"), "footer")
    watermark_text = _watermark_text(document.metadata.get("watermark"))
    context = _template_context(document)
    section = word_document.sections[0]
    content_width = preset.page.width - preset.page.margin_left - preset.page.margin_right

    if header_content or watermark_text:
        if header_content:
            _configure_page_region_word_style(
                word_document,
                "Header",
                _page_style(preset, "header"),
            )
        header = section.header
        paragraph = header.paragraphs[0]
        if header_content:
            style = _page_style(preset, "header")
            _render_page_region(paragraph, header_content, style, context, content_width)
        if watermark_text:
            _add_text_watermark(paragraph, watermark_text, preset)

    if footer_content:
        _configure_page_region_word_style(
            word_document,
            "Footer",
            _page_style(preset, "footer"),
        )
        footer = section.footer
        paragraph = footer.paragraphs[0]
        style = _page_style(preset, "footer")
        _render_page_region(paragraph, footer_content, style, context, content_width)

def _page_region_content(raw: Any, region_name: str) -> dict[str, str]:
    if raw in (None, False):
        return {}
    if isinstance(raw, str):
        if region_name == "header":
            raise ValueError(
                "Frontmatter 'header' does not support centered string content; "
                "use 'left' and/or 'right'"
            )
        return {"center": raw}
    if not isinstance(raw, Mapping):
        raise TypeError(f"Frontmatter '{region_name}' must be a mapping or string")

    allowed = (
        {"left", "right"}
        if region_name == "header"
        else {"left", "center", "right"}
    )
    unexpected = sorted(str(key) for key in raw if key not in allowed)
    if unexpected:
        names = ", ".join(unexpected)
        raise ValueError(f"Unsupported '{region_name}' fields: {names}")

    return {
        position: str(value)
        for position in ("left", "center", "right")
        if (value := raw.get(position)) not in (None, "")
    }


def _watermark_text(raw: Any) -> str:
    if raw in (None, False):
        return ""
    if isinstance(raw, str):
        return raw
    if not isinstance(raw, Mapping):
        raise TypeError("Frontmatter 'watermark' must be a mapping or string")

    unexpected = sorted(str(key) for key in raw if key not in {"text", "enabled"})
    if unexpected:
        names = ", ".join(unexpected)
        raise ValueError(f"Unsupported 'watermark' fields: {names}")
    if raw.get("enabled", True) is False:
        return ""
    value = raw.get("text", "")
    return "" if value is None else str(value)


def _template_context(document: Document) -> dict[str, str]:
    raw_date = document.metadata.get("date")
    return {
        "title": _document_title(document),
        "date": (
            str(raw_date)
            if raw_date is not None
            else datetime.now(UTC).date().isoformat()
        ),
    }


def _page_style(preset: Preset, name: str) -> Style:
    return preset.page_styles.get(name, preset.styles["body"])


def _configure_page_region_word_style(
    word_document: WordprocessingDocument,
    name: str,
    style: Style,
) -> None:
    try:
        word_style = word_document.styles[name]
    except KeyError:
        word_style = word_document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        word_style.base_style = word_document.styles["Normal"]

    word_style.paragraph_format.tab_stops.clear_all()
    _configure_word_style_font(word_style, style)

    character_style_name = f"{name} Char"
    try:
        character_style = word_document.styles[character_style_name]
    except KeyError:
        character_style = word_document.styles.add_style(
            character_style_name,
            WD_STYLE_TYPE.CHARACTER,
        )
    _configure_word_style_font(character_style, style)


def _configure_word_style_font(word_style, style: Style) -> None:
    word_style.font.name = style.font_ascii
    word_style.font.size = Pt(style.size_half_points / 2)
    word_style.font.bold = style.bold
    word_style.font.italic = style.italic

    run_properties = word_style.element.get_or_add_rPr()
    r_fonts = run_properties.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), style.font_ascii)
    r_fonts.set(qn("w:hAnsi"), style.font_ascii)
    r_fonts.set(qn("w:eastAsia"), style.font_east_asia)
    r_fonts.set(qn("w:cs"), style.font_ascii)

    complex_size = run_properties.find(qn("w:szCs"))
    if complex_size is None:
        complex_size = OxmlElement("w:szCs")
        run_properties.append(complex_size)
    complex_size.set(qn("w:val"), str(style.size_half_points))


def _render_page_region(
    paragraph,
    content: Mapping[str, str],
    style: Style,
    context: Mapping[str, str],
    content_width: int,
) -> None:
    _apply_paragraph_style(paragraph, style)
    _apply_paragraph_run_style(paragraph, style)
    positions = [
        position
        for position in ("left", "center", "right")
        if content.get(position)
    ]
    if not positions:
        return

    tab_stops = paragraph.paragraph_format.tab_stops
    tab_stops.clear_all()
    if len(positions) == 1:
        paragraph.alignment = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
        }[positions[0]]
        _append_template(paragraph, content[positions[0]], style, context)
        return

    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    has_left = "left" in positions
    has_center = "center" in positions
    has_right = "right" in positions

    if has_center:
        tab_stops.add_tab_stop(Twips(content_width / 2), WD_TAB_ALIGNMENT.CENTER)
    if has_right:
        tab_stops.add_tab_stop(Twips(content_width), WD_TAB_ALIGNMENT.RIGHT)

    if has_left:
        _append_template(paragraph, content["left"], style, context)

    if has_center:
        _add_styled_run(paragraph, "\t", style)
        _append_template(paragraph, content["center"], style, context)

    if has_right:
        _add_styled_run(paragraph, "\t", style)
        _append_template(paragraph, content["right"], style, context)


def _append_template(
    paragraph, template: str, style: Style, context: Mapping[str, str]
) -> None:
    cursor = 0
    for match in TEMPLATE_TOKEN_RE.finditer(template):
        if match.start() > cursor:
            _add_styled_run(paragraph, template[cursor : match.start()], style)

        token = match.group(1)
        field_code = PAGE_FIELD_CODES.get(token)
        if field_code:
            _add_word_field(paragraph, field_code, style)
        else:
            _add_styled_run(paragraph, context[token], style)
        cursor = match.end()

    if cursor < len(template):
        _add_styled_run(paragraph, template[cursor:], style)


def _add_word_field(paragraph, instruction: str, style: Style) -> None:
    begin_run = _add_styled_run(paragraph, "", style)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run._r.append(begin)

    instruction_run = _add_styled_run(paragraph, "", style)
    instruction_text = OxmlElement("w:instrText")
    instruction_text.set(qn("xml:space"), "preserve")
    instruction_text.text = f" {instruction} "
    instruction_run._r.append(instruction_text)

    separate_run = _add_styled_run(paragraph, "", style)
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run._r.append(separate)

    result_run = paragraph.add_run("1")
    _apply_run_style(result_run, style)

    end_run = _add_styled_run(paragraph, "", style)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def _add_text_watermark(paragraph, text: str, preset: Preset) -> None:
    style = _page_style(preset, "watermark")
    settings = preset.watermark_defaults
    color = settings.color.strip().lstrip("#").upper()
    if not re.fullmatch(r"[0-9A-F]{6}", color):
        raise ValueError("Watermark color must be a six-digit hexadecimal RGB value")
    if not 0 <= settings.opacity <= 1:
        raise ValueError("Watermark opacity must be between 0 and 1")

    pict = OxmlElement("w:pict")
    shape_type = etree.Element(
        etree.QName(VML_NS, "shapetype"),
        nsmap={"v": VML_NS, "o": OFFICE_NS},
    )
    shape_type.set("id", "_x0000_t136")
    shape_type.set("coordsize", "21600,21600")
    shape_type.set(etree.QName(OFFICE_NS, "spt"), "136")
    shape_type.set("adj", "10800")
    shape_type.set("path", "m@7,l@8,m@5,21600l@6,21600e")

    formulas = etree.SubElement(shape_type, etree.QName(VML_NS, "formulas"))
    for equation in (
        "sum #0 0 10800",
        "prod #0 2 1",
        "sum 21600 0 @1",
        "sum 0 0 @2",
        "sum 21600 0 @3",
        "if @0 @3 0",
        "if @0 21600 @1",
        "if @0 0 @2",
        "if @0 @4 21600",
        "mid @5 @6",
        "mid @8 @5",
        "mid @7 @8",
        "mid @6 @7",
        "sum @6 0 @5",
    ):
        formula = etree.SubElement(formulas, etree.QName(VML_NS, "f"))
        formula.set("eqn", equation)

    shape_path = etree.SubElement(shape_type, etree.QName(VML_NS, "path"))
    shape_path.set("textpathok", "t")
    shape_path.set(etree.QName(OFFICE_NS, "connecttype"), "custom")
    shape_path.set(
        etree.QName(OFFICE_NS, "connectlocs"),
        "@9,0;@10,10800;@11,21600;@12,10800",
    )
    shape_path.set(
        etree.QName(OFFICE_NS, "connectangles"), "270,180,90,0"
    )
    shape_type_textpath = etree.SubElement(
        shape_type, etree.QName(VML_NS, "textpath")
    )
    shape_type_textpath.set("on", "t")
    shape_type_textpath.set("fitshape", "t")
    handles = etree.SubElement(shape_type, etree.QName(VML_NS, "handles"))
    handle = etree.SubElement(handles, etree.QName(VML_NS, "h"))
    handle.set("position", "#0,bottomRight")
    handle.set("xrange", "6629,14971")
    lock = etree.SubElement(shape_type, etree.QName(OFFICE_NS, "lock"))
    lock.set(etree.QName(VML_NS, "ext"), "edit")
    lock.set("text", "t")
    lock.set("shapetype", "t")

    shape = etree.Element(
        etree.QName(VML_NS, "shape"),
        nsmap={"v": VML_NS, "o": OFFICE_NS, "w10": WORD_2003_NS},
    )
    shape.set("id", "PowerPlusWaterMarkObject1")
    shape.set(etree.QName(OFFICE_NS, "spid"), "_x0000_s2049")
    shape.set("type", "#_x0000_t136")
    shape.set(etree.QName(OFFICE_NS, "allowincell"), "f")
    shape.set("fillcolor", f"#{color}")
    shape.set("stroked", "f")
    rotation = settings.rotation % 360
    shape.set(
        "style",
        ";".join(
            (
                "position:absolute",
                "margin-left:0",
                "margin-top:0",
                "width:500pt",
                "height:165pt",
                f"rotation:{rotation}",
                "z-index:-251654144",
                "mso-position-horizontal:center",
                "mso-position-horizontal-relative:margin",
                "mso-position-vertical:center",
                "mso-position-vertical-relative:margin",
            )
        ),
    )

    fill = etree.SubElement(shape, etree.QName(VML_NS, "fill"))
    fill.set("opacity", f"{settings.opacity:g}")
    textpath = etree.SubElement(shape, etree.QName(VML_NS, "textpath"))
    textpath.set("string", text)
    font_style = (
        f'font-family:"{style.font_east_asia}";'
        f"font-size:{style.size_half_points / 2:g}pt;"
        "v-text-align:center"
    )
    if style.bold:
        font_style += ";font-weight:bold"
    textpath.set("style", font_style)
    wrap = etree.SubElement(shape, etree.QName(WORD_2003_NS, "wrap"))
    wrap.set("anchorx", "margin")
    wrap.set("anchory", "margin")

    pict.append(shape_type)
    pict.append(shape)
    run = paragraph.add_run()
    run._r.append(pict)


def _set_core_properties(word_document: WordprocessingDocument, title: str) -> None:
    properties = word_document.core_properties
    properties.title = title
    properties.author = "mdstyledocx"
    properties.last_modified_by = "mdstyledocx"
    now = datetime.now(UTC).replace(microsecond=0, tzinfo=None)
    properties.created = now
    properties.modified = now


def _document_title(document: Document) -> str:
    if document.metadata.get("title"):
        return str(document.metadata["title"])
    for block in document.blocks:
        if block.kind == "heading" and block.level == 1:
            return "".join(span.text for span in block.spans if isinstance(span, InlineSpan)).strip()
    return "Document"


def _append_block(
    word_document: WordprocessingDocument,
    block: Block,
    state: BuildState,
    *,
    next_block: Block | None = None,
) -> None:
    if block.kind == "page_break":
        word_document.add_page_break()
        return

    if block.kind == "table":
        _append_table(word_document, block, state)
        return

    rendered_spans = _rendered_spans(block, state)
    style = _resolve_style(block, state.preset)
    if (
        block.kind == "heading"
        and block.level == 1
        and next_block is not None
        and next_block.kind == "date"
    ):
        style = replace(style, spacing_after=0)
    if _spans_have_image(rendered_spans):
        style = replace(style, line=240, line_rule="auto")

    paragraph = word_document.add_paragraph()
    _apply_paragraph_style(paragraph, style)

    if block.kind == "list_item":
        prefix = "• " if block.list_kind == "bullet" else f"{block.number}. "
        _add_text_run(paragraph, InlineSpan(text=prefix), style)

    for span in rendered_spans:
        if isinstance(span, ImageSpan):
            _add_image_run(paragraph, span, state)
        elif span.text:
            _add_text_run(paragraph, span, style)

    if not paragraph.runs:
        paragraph.add_run("")


def _append_table(
    word_document: WordprocessingDocument, block: Block, state: BuildState
) -> None:
    if not block.table_rows:
        return

    column_count = len(block.table_rows[0])
    if not column_count or any(len(row) != column_count for row in block.table_rows):
        raise ValueError("Table rows must have a consistent, non-zero column count")

    table = word_document.add_table(rows=len(block.table_rows), cols=column_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    column_widths = _content_aware_column_widths(block, state.preset)
    _set_table_width(table, sum(column_widths))
    for column_index, width in enumerate(column_widths):
        table.columns[column_index].width = Twips(width)

    body_style = replace(
        state.preset.styles["body"],
        first_line_indent=0,
        left_indent=0,
        hanging=0,
    )

    for row_index, source_row in enumerate(block.table_rows):
        row = table.rows[row_index]
        _configure_table_row(row, is_header=row_index == 0)
        for column_index, spans in enumerate(source_row):
            cell = row.cells[column_index]
            cell.width = Twips(column_widths[column_index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            alignment = block.table_alignments[column_index]
            cell_style = replace(
                body_style,
                align=alignment or ("center" if row_index == 0 else "left"),
                bold=body_style.bold or row_index == 0,
            )
            _populate_table_cell(cell, spans, cell_style)


def _populate_table_cell(
    cell, spans: list[InlineElement], style: Style
) -> None:
    paragraph = cell.paragraphs[0]
    _apply_paragraph_style(paragraph, style)
    for span in spans:
        if isinstance(span, ImageSpan):
            raise ValueError("Images inside Markdown table cells are not supported")
        if span.text:
            _add_text_run(paragraph, span, style)
    if not paragraph.runs:
        paragraph.add_run("")


def _configure_table_row(row, *, is_header: bool) -> None:
    properties = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    properties.append(cant_split)
    if is_header:
        repeat_header = OxmlElement("w:tblHeader")
        repeat_header.set(qn("w:val"), "true")
        properties.append(repeat_header)


def _set_table_width(table, width: int) -> None:
    properties = table._tbl.tblPr
    table_width = properties.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        properties.insert(0, table_width)
    table_width.set(qn("w:type"), "dxa")
    table_width.set(qn("w:w"), str(width))


def _content_aware_column_widths(block: Block, preset: Preset) -> list[int]:
    column_count = len(block.table_rows[0])
    available_width = (
        preset.page.width - preset.page.margin_left - preset.page.margin_right
    )
    minimum_width = min(720, max(240, available_width // (column_count * 2)))

    weights: list[int] = []
    for column_index in range(column_count):
        natural_width = max(
            _display_width(_inline_text(row[column_index]))
            for row in block.table_rows
        )
        weights.append(max(4, min(60, natural_width)))

    distributable_width = max(0, available_width - minimum_width * column_count)
    total_weight = sum(weights)
    widths = [
        minimum_width + distributable_width * weight // total_weight
        for weight in weights
    ]
    widths[-1] += available_width - sum(widths)
    return widths


def _inline_text(spans: list[InlineElement]) -> str:
    return "".join(
        span.text if isinstance(span, InlineSpan) else span.alt_text
        for span in spans
    )


def _display_width(value: str) -> int:
    line_widths = [
        sum(
            2 if unicodedata.east_asian_width(character) in {"W", "F"} else 1
            for character in line
        )
        for line in value.splitlines() or [""]
    ]
    return max(line_widths, default=0)


def _resolve_style(block: Block, preset: Preset) -> Style:
    if block.kind == "heading":
        key = {1: "title", 2: "heading1", 3: "heading2"}.get(block.level, "heading3")
        return preset.styles[key]

    base = preset.styles["body"]
    if block.kind == "date":
        return replace(
            base,
            align="center",
            first_line_indent=0,
            left_indent=0,
            hanging=0,
            spacing_after=preset.styles["title"].spacing_after,
        )

    if block.kind == "list_item":
        left_indent = (
            preset.list_settings.base_left_indent
            + block.list_level * preset.list_settings.level_step
        )
        return replace(
            base,
            first_line_indent=0,
            left_indent=left_indent,
            hanging=preset.list_settings.hanging,
        )

    return base


def _apply_paragraph_style(paragraph, style: Style) -> None:
    alignment_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "both": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    if style.align:
        paragraph.alignment = alignment_map[style.align]

    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Twips(style.spacing_before)
    paragraph_format.space_after = Twips(style.spacing_after)
    paragraph_format.left_indent = Twips(style.left_indent)

    if style.hanging:
        paragraph_format.first_line_indent = Twips(-style.hanging)
    else:
        paragraph_format.first_line_indent = Twips(style.first_line_indent)

    if style.line_rule == "exact":
        paragraph_format.line_spacing = Pt(style.line / 20)
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    elif style.line_rule == "auto":
        paragraph_format.line_spacing = style.line / 240
    else:
        paragraph_format.line_spacing = style.line / 240

    _apply_overflow_punctuation(paragraph, style.overflow_punctuation)
    _apply_bottom_border(paragraph, style)


def _apply_overflow_punctuation(paragraph, value: bool | None) -> None:
    if value is None:
        return

    properties = paragraph._p.get_or_add_pPr()
    overflow = properties.find(qn("w:overflowPunct"))
    if overflow is None:
        overflow = OxmlElement("w:overflowPunct")
        spacing = properties.find(qn("w:spacing"))
        if spacing is None:
            properties.append(overflow)
        else:
            properties.insert(properties.index(spacing), overflow)
    overflow.set(qn("w:val"), "1" if value else "0")


def _apply_bottom_border(paragraph, style: Style) -> None:
    if not style.bottom_border_size:
        return

    properties = paragraph._p.get_or_add_pPr()
    borders = properties.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        properties.insert_element_before(
            borders,
            "w:shd",
            "w:tabs",
            "w:spacing",
            "w:ind",
            "w:jc",
            "w:rPr",
            "w:sectPr",
        )

    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(style.bottom_border_size))
    bottom.set(qn("w:space"), str(style.bottom_border_space))
    bottom.set(qn("w:color"), style.bottom_border_color)


def _add_text_run(paragraph, span: InlineSpan, style: Style) -> None:
    run = paragraph.add_run(span.text)
    font_ascii = style.font_ascii
    font_east_asia = style.font_east_asia

    if span.code:
        font_ascii = "Consolas"
        font_east_asia = "等线"

    _apply_run_style(
        run,
        style,
        font_ascii=font_ascii,
        font_east_asia=font_east_asia,
        bold=style.bold or span.bold,
        italic=style.italic or span.italic,
    )


def _add_styled_run(paragraph, text: str, style: Style):
    run = paragraph.add_run(text)
    _apply_run_style(run, style)
    return run


def _apply_run_style(
    run,
    style: Style,
    *,
    font_ascii: str | None = None,
    font_east_asia: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    font_ascii = font_ascii or style.font_ascii
    font_east_asia = font_east_asia or style.font_east_asia
    run.font.name = font_ascii
    run.font.size = Pt(style.size_half_points / 2)
    run.bold = style.bold if bold is None else bold
    run.italic = style.italic if italic is None else italic

    r_fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), font_ascii)
    r_fonts.set(qn("w:hAnsi"), font_ascii)
    r_fonts.set(qn("w:eastAsia"), font_east_asia)
    r_fonts.set(qn("w:cs"), font_ascii)

    run_properties = run._element.get_or_add_rPr()
    complex_size = run_properties.find(qn("w:szCs"))
    if complex_size is None:
        complex_size = OxmlElement("w:szCs")
        run_properties.append(complex_size)
    complex_size.set(qn("w:val"), str(style.size_half_points))


def _apply_paragraph_run_style(paragraph, style: Style) -> None:
    properties = paragraph._p.get_or_add_pPr()
    run_properties = properties.find(qn("w:rPr"))
    if run_properties is None:
        run_properties = OxmlElement("w:rPr")
        properties.insert_element_before(run_properties, "w:sectPr")

    r_fonts = run_properties.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        run_properties.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), style.font_ascii)
    r_fonts.set(qn("w:hAnsi"), style.font_ascii)
    r_fonts.set(qn("w:eastAsia"), style.font_east_asia)
    r_fonts.set(qn("w:cs"), style.font_ascii)

    for tag in ("w:sz", "w:szCs"):
        size = run_properties.find(qn(tag))
        if size is None:
            size = OxmlElement(tag)
            run_properties.append(size)
        size.set(qn("w:val"), str(style.size_half_points))


def _add_image_run(paragraph, span: ImageSpan, state: BuildState) -> None:
    run = paragraph.add_run()
    inline_shape = run.add_picture(span.path)
    max_width = _max_image_width(state.preset)
    if inline_shape.width > max_width:
        scale = max_width / inline_shape.width
        inline_shape.width = max_width
        inline_shape.height = int(inline_shape.height * scale)

    description = span.alt_text or os.path.basename(span.path)
    inline_shape._inline.docPr.set("descr", description)
    inline_shape._inline.docPr.set("name", description)


def _max_image_width(preset: Preset) -> int:
    return Twips(preset.page.width - preset.page.margin_left - preset.page.margin_right)


def _spans_have_image(spans: list[InlineElement]) -> bool:
    return any(isinstance(span, ImageSpan) for span in spans)


def _rendered_spans(block: Block, state: BuildState) -> list[InlineElement]:
    if block.kind != "heading":
        return block.spans

    scheme = state.preset.heading_numbering.get(block.level)
    if not scheme:
        return block.spans

    _advance_heading_counters(state, block.level)
    if _has_number_prefix(block.spans, scheme):
        return block.spans

    prefix = _format_heading_prefix(scheme, state.heading_counters[block.level])
    return [InlineSpan(text=prefix)] + block.spans


def _advance_heading_counters(state: BuildState, level: int) -> None:
    state.heading_counters[level] = state.heading_counters.get(level, 0) + 1
    for deeper_level in list(state.heading_counters):
        if deeper_level > level:
            state.heading_counters[deeper_level] = 0


def _has_number_prefix(spans: list[InlineElement], scheme: str) -> bool:
    text = "".join(span.text for span in spans if isinstance(span, InlineSpan)).lstrip()
    patterns = {
        "cn-section": r"^[一二三四五六七八九十百千万零〇两]+、",
        "cn-paren": r"^（[一二三四五六七八九十百千万零〇两]+）",
        "arabic-dot": r"^\d+[.．]\s*",
    }
    return re.match(patterns[scheme], text) is not None


def _format_heading_prefix(scheme: str, number: int) -> str:
    if scheme == "cn-section":
        return f"{_to_chinese_number(number)}、"
    if scheme == "cn-paren":
        return f"（{_to_chinese_number(number)}）"
    if scheme == "arabic-dot":
        return f"{number}. "
    raise ValueError(f"Unsupported heading numbering scheme: {scheme}")


def _to_chinese_number(number: int) -> str:
    digits = "零一二三四五六七八九"
    units = ["", "十", "百", "千"]
    raw = str(number)
    parts: list[str] = []

    for index, char in enumerate(raw):
        digit = int(char)
        unit_index = len(raw) - index - 1
        if digit == 0:
            if parts and parts[-1] != "零" and any(next_char != "0" for next_char in raw[index + 1 :]):
                parts.append("零")
            continue
        if not (digit == 1 and unit_index == 1 and not parts and len(raw) == 2):
            parts.append(digits[digit])
        parts.append(units[unit_index])

    return "".join(parts) or "零"
