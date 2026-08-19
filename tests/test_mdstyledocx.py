from __future__ import annotations

import base64
import io
import json
import re
import tempfile
import unittest
import zipfile
from contextlib import redirect_stdout
from pathlib import Path

from docx import Document as WordDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from mdstyledocx import __version__
from mdstyledocx.cli import main
from mdstyledocx.docx_writer import build_docx
from mdstyledocx.markdown import parse_markdown
from mdstyledocx.presets import (
    list_presets,
    load_preset,
    load_preset_definition,
    load_preset_rules,
    load_preset_schema,
)

SAMPLE_MARKDOWN = """# 关于开展示例工作的通知

各有关单位：

为统一输出格式，现将有关事项通知如下。

## 一、工作目标

1. 统一内容源。
2. 统一输出格式。

<!-- pagebreak -->

## 二、工作要求

请各单位按要求执行。
"""

SAMPLE_MARKDOWN_WITH_SUBHEADING = """# 关于开展示例工作的通知

各有关单位：

为统一输出格式，现将有关事项通知如下。

## 一、工作目标

### （一）总体要求

请各单位按要求执行。
"""

SAMPLE_MARKDOWN_WITH_UNNUMBERED_HEADINGS = """# 关于开展示例工作的通知

各有关单位：

## 工作目标

### 总体要求

#### 任务分工

请各单位按要求执行。
"""

SAMPLE_MARKDOWN_WITH_MIXED_HEADING_NUMBERS = """# 关于开展示例工作的通知

## 一、工作目标

### （一）总体要求

#### 1. 任务分工

## 结语

请各单位按要求执行。
"""

SAMPLE_MARKDOWN_WITH_PAGE_CONTENT = """---
title: 关于开展示例工作的通知
date: 2026-08-19
header:
  left: "某某单位 {title}"
  right: 内部材料
footer:
  left: "{date}"
  center: "— {page} / {pages} —"
  right: 校对稿
watermark:
  text: 内部资料
---

正文内容。
"""

SAMPLE_MARKDOWN_WITH_TWO_ZONE_HEADER = """---
header:
  left: 杭州红帆智药科技有限公司
  right: 内部资料，请勿外传
footer:
  center: "— {page} —"
---

# 公司制度

正文内容。
"""

SAMPLE_MARKDOWN_WITH_TABLE = """# 项目情况表

| 序号 | 事项说明 |
| :---: | --- |
| 1 | 简要事项 |
| 2 | 这是用于验证内容较多的列能够获得更多页面宽度的较长说明，避免所有列默认等宽。 |
| 3 | A \\| B |
"""

PNG_1X1_BASE64 = (
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+yF9kAAAAASUVORK5CYII="
)


class MarkdownParsingTests(unittest.TestCase):
    def test_parse_markdown_blocks(self) -> None:
        document = parse_markdown(SAMPLE_MARKDOWN)

        self.assertEqual(document.blocks[0].kind, "heading")
        self.assertEqual(document.blocks[0].level, 1)
        self.assertEqual(document.blocks[1].kind, "paragraph")
        self.assertEqual(document.blocks[3].kind, "heading")
        self.assertEqual(document.blocks[4].kind, "list_item")
        self.assertEqual(document.blocks[6].kind, "page_break")

    def test_parse_markdown_image_span_uses_base_path(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            image_path = temp_path / "seal.png"
            image_path.write_bytes(base64.b64decode(PNG_1X1_BASE64))

            document = parse_markdown("![公章](seal.png)", base_path=temp_path)

            image_span = document.blocks[0].spans[0]
            self.assertEqual(image_span.alt_text, "公章")
            self.assertEqual(image_span.path, str(image_path.resolve()))

    def test_parse_markdown_table_with_alignment_and_escaped_pipe(self) -> None:
        document = parse_markdown(SAMPLE_MARKDOWN_WITH_TABLE)

        table = document.blocks[1]
        self.assertEqual(table.kind, "table")
        self.assertEqual(table.table_alignments, ["center", None])
        self.assertEqual(len(table.table_rows), 4)
        self.assertEqual(table.table_rows[0][0][0].text, "序号")
        self.assertEqual(table.table_rows[3][1][0].text, "A | B")

    def test_parse_nested_yaml_frontmatter(self) -> None:
        document = parse_markdown(SAMPLE_MARKDOWN_WITH_PAGE_CONTENT)

        self.assertEqual(
            document.metadata["header"]["left"],
            "某某单位 {title}",
        )
        self.assertEqual(document.metadata["footer"]["center"], "— {page} / {pages} —")
        self.assertEqual(document.metadata["watermark"]["text"], "内部资料")
        self.assertEqual(document.blocks[0].kind, "heading")
        self.assertEqual(document.blocks[0].level, 1)
        self.assertEqual(document.blocks[1].kind, "date")
        self.assertEqual(document.blocks[1].spans[0].text, "2026-08-19")

    def test_invalid_yaml_frontmatter_is_rejected(self) -> None:
        with self.assertRaisesRegex(ValueError, "Invalid YAML frontmatter"):
            parse_markdown("---\nheader: [\n---\n正文")

    def test_frontmatter_requires_top_level_mapping(self) -> None:
        with self.assertRaisesRegex(TypeError, "mapping at the top level"):
            parse_markdown("---\n- first\n- second\n---\n正文")

    def test_header_center_zone_is_rejected(self) -> None:
        document = parse_markdown(
            "---\nheader:\n  center: 不支持的中页眉\n---\n\n# 标题"
        )

        with self.assertRaisesRegex(ValueError, "Unsupported 'header' fields: center"):
            build_docx(document, load_preset("official-doc-cn-system-fonts-12pt"))


class PresetLoadingTests(unittest.TestCase):
    def test_preset_schema_is_packaged_and_versioned(self) -> None:
        schema = json.loads(load_preset_schema())

        self.assertEqual(schema["$schema"], "https://json-schema.org/draft/2020-12/schema")
        self.assertEqual(schema["properties"]["schema_version"]["const"], 1)
        self.assertIn("style", schema["$defs"])
        self.assertIn("overflow_punctuation", schema["$defs"]["style"]["properties"])

    def test_builtin_presets_are_loaded_from_spec_files(self) -> None:
        presets = dict(list_presets())

        self.assertIn("default", presets)
        self.assertIn("official-doc-cn", presets)
        self.assertIn("official-doc-cn-12pt", presets)
        self.assertIn("official-doc-cn-system-fonts", presets)
        self.assertIn("official-doc-cn-system-fonts-12pt", presets)
        self.assertNotIn("gov-cn", presets)
        self.assertNotIn("gov-cn-hei", presets)
        self.assertEqual(
            presets["official-doc-cn"], "Chinese official document baseline preset"
        )

    def test_preset_rules_are_available(self) -> None:
        rules = load_preset_rules("official-doc-cn")

        self.assertIn("推荐 Markdown 写法", rules)
        self.assertIn("仿宋_GB2312", rules)

    def test_legacy_preset_names_are_compatibility_aliases(self) -> None:
        self.assertEqual(load_preset("gov-cn").name, "official-doc-cn")
        self.assertEqual(
            load_preset("gov-cn-hei").name, "official-doc-cn-system-fonts"
        )
        self.assertEqual(
            load_preset_rules("gov-cn"), load_preset_rules("official-doc-cn")
        )

    def test_official_doc_presets_apply_global_paragraph_defaults(self) -> None:
        for name in ("official-doc-cn", "official-doc-cn-system-fonts"):
            with self.subTest(preset=name):
                preset = load_preset(name)

                self.assertEqual(preset.styles["title"].size_half_points, 44)
                for style_name in ("heading1", "heading2", "heading3", "body"):
                    self.assertEqual(preset.styles[style_name].size_half_points, 32)
                for style in preset.styles.values():
                    self.assertEqual(style.line, 560)
                    self.assertEqual(style.line_rule, "exact")
                    self.assertEqual(style.spacing_before, 0)
                self.assertEqual(preset.styles["title"].spacing_after, 560)
                for style_name in ("heading1", "heading2", "heading3", "body"):
                    self.assertEqual(preset.styles[style_name].spacing_after, 0)
                self.assertEqual(preset.styles["heading1"].first_line_indent, 640)
                self.assertEqual(preset.styles["heading2"].first_line_indent, 640)
                self.assertEqual(preset.styles["body"].align, "both")
                self.assertFalse(preset.styles["body"].overflow_punctuation)

    def test_12pt_presets_apply_compact_hierarchy_and_global_spacing(self) -> None:
        for name in ("official-doc-cn-12pt", "official-doc-cn-system-fonts-12pt"):
            with self.subTest(preset=name):
                preset = load_preset(name)

                self.assertEqual(preset.styles["title"].size_half_points, 36)
                for style_name in ("heading1", "heading2", "heading3", "body"):
                    self.assertEqual(preset.styles[style_name].size_half_points, 24)
                self.assertEqual(preset.styles["body"].first_line_indent, 480)
                for style in preset.styles.values():
                    self.assertEqual(style.line, 400)
                    self.assertEqual(style.line_rule, "exact")
                    self.assertEqual(style.spacing_before, 0)
                self.assertEqual(preset.styles["title"].spacing_after, 400)
                for style_name in ("heading1", "heading2", "heading3", "body"):
                    self.assertEqual(preset.styles[style_name].spacing_after, 0)
                self.assertEqual(preset.styles["heading1"].first_line_indent, 480)
                self.assertEqual(preset.styles["heading2"].first_line_indent, 480)
                self.assertEqual(preset.styles["body"].align, "both")
                self.assertFalse(preset.styles["body"].overflow_punctuation)

    def test_official_doc_presets_leave_heading_numbering_to_markdown(self) -> None:
        for name in (
            "official-doc-cn",
            "official-doc-cn-12pt",
            "official-doc-cn-system-fonts",
            "official-doc-cn-system-fonts-12pt",
        ):
            with self.subTest(preset=name):
                self.assertEqual(load_preset(name).heading_numbering, {})

    def test_presets_define_page_content_styles(self) -> None:
        for name in (
            "default",
            "official-doc-cn",
            "official-doc-cn-12pt",
            "official-doc-cn-system-fonts",
            "official-doc-cn-system-fonts-12pt",
        ):
            with self.subTest(preset=name):
                preset = load_preset(name)

                self.assertIn("header", preset.page_styles)
                self.assertIn("footer", preset.page_styles)
                self.assertIn("watermark", preset.page_styles)
                self.assertEqual(
                    preset.page_styles["header"].size_half_points,
                    preset.page_styles["footer"].size_half_points,
                )
                if name.startswith("official-doc-cn"):
                    self.assertEqual(
                        preset.page_styles["header"].size_half_points,
                        21,
                    )
                self.assertEqual(preset.watermark_defaults.color, "D9D9D9")
                self.assertEqual(preset.watermark_defaults.opacity, 0.25)
                self.assertEqual(preset.watermark_defaults.rotation, -45)

    def test_style_override_takes_priority_over_paragraph_defaults(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            preset_file = Path(temp_dir) / "override.json"
            preset_file.write_text(
                json.dumps(
                    {
                        "extends": "official-doc-cn",
                        "paragraph_defaults": {"line": 520},
                        "styles": {"title": {"line": 600}},
                    }
                ),
                encoding="utf-8",
            )

            preset = load_preset("default", preset_file)

        self.assertEqual(preset.styles["body"].line, 520)
        self.assertEqual(preset.styles["heading1"].line, 520)
        self.assertEqual(preset.styles["title"].line, 600)

    def test_resolved_preset_definition_applies_inheritance(self) -> None:
        definition = load_preset_definition("official-doc-cn-12pt")

        self.assertNotIn("extends", definition)
        self.assertEqual(definition["schema_version"], 1)
        self.assertEqual(definition["paragraph_defaults"]["line"], 400)
        self.assertEqual(definition["styles"]["title"]["size_half_points"], 36)
        self.assertIn("watermark", definition["page_styles"])

    def test_custom_preset_rejects_unknown_schema_or_fields(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            unsupported_version = temp_path / "version.json"
            unsupported_version.write_text(
                json.dumps({"schema_version": 2, "extends": "default"}),
                encoding="utf-8",
            )
            unknown_field = temp_path / "field.json"
            unknown_field.write_text(
                json.dumps(
                    {
                        "schema_version": 1,
                        "extends": "default",
                        "line_spacing": 400,
                    }
                ),
                encoding="utf-8",
            )

            with self.assertRaisesRegex(ValueError, "Unsupported preset schema_version"):
                load_preset("default", unsupported_version)
            with self.assertRaisesRegex(ValueError, "Unsupported preset fields"):
                load_preset("default", unknown_field)


class DocxGenerationTests(unittest.TestCase):
    def test_official_doc_preset_contains_expected_markers(self) -> None:
        document = parse_markdown(SAMPLE_MARKDOWN)
        payload = build_docx(document, load_preset("official-doc-cn"))

        with zipfile.ZipFile(io.BytesIO(payload)) as archive:
            names = set(archive.namelist())
            xml = archive.read("word/document.xml").decode("utf-8")

        self.assertIn("word/document.xml", names)
        self.assertIn("docProps/core.xml", names)
        self.assertIn("仿宋_GB2312", xml)
        self.assertIn("方正小标宋简体", xml)
        self.assertIn('w:sz w:val="44"', xml)
        self.assertIn('w:sz w:val="32"', xml)
        self.assertIn('w:firstLine="640"', xml)
        self.assertIn('w:line="560"', xml)
        self.assertIn('w:lineRule="exact"', xml)
        self.assertNotIn("一、一、工作目标", xml)

    def test_system_fonts_preset_contains_expected_markers(self) -> None:
        document = parse_markdown(SAMPLE_MARKDOWN_WITH_SUBHEADING)
        payload = build_docx(document, load_preset("official-doc-cn-system-fonts"))

        with zipfile.ZipFile(io.BytesIO(payload)) as archive:
            xml = archive.read("word/document.xml").decode("utf-8")

        self.assertIn("黑体", xml)
        self.assertIn("楷体", xml)
        self.assertIn("仿宋", xml)
        self.assertNotIn("方正小标宋简体", xml)
        self.assertNotIn("楷体_GB2312", xml)
        self.assertNotIn("仿宋_GB2312", xml)

    def test_official_doc_presets_preserve_mixed_heading_numbers(self) -> None:
        document = parse_markdown(SAMPLE_MARKDOWN_WITH_MIXED_HEADING_NUMBERS)
        payload = build_docx(document, load_preset("official-doc-cn-system-fonts"))

        with zipfile.ZipFile(io.BytesIO(payload)) as archive:
            xml = archive.read("word/document.xml").decode("utf-8")
        visible_text = re.sub(r"<[^>]+>", "", xml)

        self.assertIn("一、工作目标", visible_text)
        self.assertIn("（一）总体要求", visible_text)
        self.assertIn("1. 任务分工", visible_text)
        self.assertIn("结语", visible_text)
        self.assertNotIn("二、结语", visible_text)

    def test_custom_preset_can_opt_in_to_automatic_heading_numbering(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            preset_file = Path(temp_dir) / "numbered.json"
            preset_file.write_text(
                json.dumps(
                    {
                        "schema_version": 1,
                        "extends": "official-doc-cn-system-fonts",
                        "heading_numbering": {
                            "2": "cn-section",
                            "3": "cn-paren",
                            "4": "arabic-dot",
                        },
                    }
                ),
                encoding="utf-8",
            )
            document = parse_markdown(SAMPLE_MARKDOWN_WITH_UNNUMBERED_HEADINGS)
            payload = build_docx(document, load_preset("default", preset_file))

        with zipfile.ZipFile(io.BytesIO(payload)) as archive:
            xml = archive.read("word/document.xml").decode("utf-8")
        visible_text = re.sub(r"<[^>]+>", "", xml)

        self.assertIn("一、工作目标", visible_text)
        self.assertIn("（一）总体要求", visible_text)
        self.assertIn("1. 任务分工", visible_text)

    def test_12pt_preset_uses_compact_hierarchy_and_spacing(self) -> None:
        document = parse_markdown(SAMPLE_MARKDOWN_WITH_UNNUMBERED_HEADINGS)
        payload = build_docx(document, load_preset("official-doc-cn-12pt"))

        with zipfile.ZipFile(io.BytesIO(payload)) as archive:
            xml = archive.read("word/document.xml").decode("utf-8")

        self.assertIn('w:sz w:val="36"', xml)
        self.assertIn('w:sz w:val="24"', xml)
        self.assertIn('w:firstLine="480"', xml)
        self.assertIn('w:line="400"', xml)
        self.assertIn('w:lineRule="exact"', xml)
        self.assertNotIn('w:sz w:val="32"', xml)

    def test_official_doc_layout_uses_title_gap_heading_indents_and_justified_body(
        self,
    ) -> None:
        document = parse_markdown(SAMPLE_MARKDOWN_WITH_SUBHEADING)
        payload = build_docx(
            document, load_preset("official-doc-cn-system-fonts-12pt")
        )
        word_document = WordDocument(io.BytesIO(payload))

        title = word_document.paragraphs[0]
        heading1 = next(
            p for p in word_document.paragraphs if p.text == "一、工作目标"
        )
        heading2 = next(
            p for p in word_document.paragraphs if p.text == "（一）总体要求"
        )
        body = next(
            p for p in word_document.paragraphs if p.text == "请各单位按要求执行。"
        )

        self.assertEqual(title.paragraph_format.space_after.twips, 400)
        self.assertEqual(heading1.paragraph_format.first_line_indent.twips, 480)
        self.assertEqual(heading2.paragraph_format.first_line_indent.twips, 480)
        self.assertEqual(body.alignment, WD_ALIGN_PARAGRAPH.JUSTIFY)
        overflow = body._p.pPr.find(qn("w:overflowPunct"))
        self.assertIsNotNone(overflow)
        self.assertEqual(overflow.get(qn("w:val")), "0")

    def test_frontmatter_date_is_centered_below_title_using_body_font(self) -> None:
        document = parse_markdown(
            "---\ndate: 2026年8月20日\n---\n\n# 公司制度\n\n正文内容。"
        )
        payload = build_docx(
            document, load_preset("official-doc-cn-system-fonts-12pt")
        )
        word_document = WordDocument(io.BytesIO(payload))

        title, date, body = word_document.paragraphs
        self.assertEqual(title.text, "公司制度")
        self.assertEqual(date.text, "2026年8月20日")
        self.assertEqual(body.text, "正文内容。")
        self.assertEqual(title.paragraph_format.space_after.twips, 0)
        self.assertEqual(date.alignment, WD_ALIGN_PARAGRAPH.CENTER)
        self.assertEqual(date.paragraph_format.first_line_indent.twips, 0)
        self.assertEqual(date.paragraph_format.space_after.twips, 400)
        self.assertEqual(date.runs[0].font.size.pt, 12)
        self.assertEqual(
            date.runs[0]._r.rPr.rFonts.get(qn("w:eastAsia")),
            "仿宋",
        )

    def test_markdown_table_has_bold_header_and_content_aware_autofit_widths(
        self,
    ) -> None:
        document = parse_markdown(SAMPLE_MARKDOWN_WITH_TABLE)
        payload = build_docx(
            document, load_preset("official-doc-cn-system-fonts-12pt")
        )
        word_document = WordDocument(io.BytesIO(payload))

        self.assertEqual(len(word_document.tables), 1)
        table = word_document.tables[0]
        self.assertTrue(table.autofit)
        self.assertEqual(table.rows[3].cells[1].text, "A | B")
        self.assertTrue(
            all(
                run.bold is True
                for cell in table.rows[0].cells
                for run in cell.paragraphs[0].runs
                if run.text
            )
        )
        self.assertTrue(
            all(
                run.bold is not True
                for cell in table.rows[1].cells
                for run in cell.paragraphs[0].runs
                if run.text
            )
        )
        self.assertLess(table.columns[0].width, table.columns[1].width)

        with zipfile.ZipFile(io.BytesIO(payload)) as archive:
            xml = archive.read("word/document.xml").decode("utf-8")
        self.assertIn('<w:tblLayout w:type="autofit"', xml)
        self.assertIn('<w:tblHeader w:val="true"', xml)

    def test_cli_writes_docx_file(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            input_path = temp_path / "notice.md"
            output_path = temp_path / "notice.docx"
            input_path.write_text(SAMPLE_MARKDOWN, encoding="utf-8")

            exit_code = main(
                [
                    str(input_path),
                    "-o",
                    str(output_path),
                    "--preset",
                    "official-doc-cn",
                ]
            )

            self.assertEqual(exit_code, 0)
            self.assertTrue(output_path.exists())
            self.assertGreater(output_path.stat().st_size, 0)

    def test_image_is_embedded_and_uses_single_spacing(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            image_path = temp_path / "seal.png"
            image_path.write_bytes(base64.b64decode(PNG_1X1_BASE64))
            markdown = "# 标题\n\n![公章](seal.png)\n"

            document = parse_markdown(markdown, base_path=temp_path)
            payload = build_docx(document, load_preset("official-doc-cn-system-fonts"))

            with zipfile.ZipFile(io.BytesIO(payload)) as archive:
                xml = archive.read("word/document.xml").decode("utf-8")
                rels = archive.read("word/_rels/document.xml.rels").decode("utf-8")
                media = set(archive.namelist())

            self.assertIn("<w:drawing>", xml)
            self.assertIn('w:line="240"', xml)
            self.assertIn('w:lineRule="auto"', xml)
            self.assertIn('relationships/image', rels)
            self.assertIn("media/image1.png", rels)
            self.assertIn("word/media/image1.png", media)

    def test_yaml_frontmatter_adds_header_footer_fields_and_watermark(self) -> None:
        document = parse_markdown(SAMPLE_MARKDOWN_WITH_PAGE_CONTENT)
        payload = build_docx(document, load_preset("official-doc-cn"))

        with zipfile.ZipFile(io.BytesIO(payload)) as archive:
            names = set(archive.namelist())
            header_xml = archive.read("word/header1.xml").decode("utf-8")
            footer_xml = archive.read("word/footer1.xml").decode("utf-8")
            settings_xml = archive.read("word/settings.xml").decode("utf-8")

        self.assertIn("word/header1.xml", names)
        self.assertIn("word/footer1.xml", names)
        self.assertIn("某某单位", header_xml)
        self.assertIn("关于开展示例工作的通知", header_xml)
        self.assertIn("内部材料", header_xml)
        self.assertIn('<v:shapetype id="_x0000_t136"', header_xml)
        self.assertIn("<v:shape", header_xml)
        self.assertIn('type="#_x0000_t136"', header_xml)
        self.assertIn('string="内部资料"', header_xml)
        self.assertIn('fillcolor="#D9D9D9"', header_xml)
        self.assertIn("rotation:315", header_xml)
        self.assertIn("<w10:wrap", header_xml)
        self.assertIn("<w:instrText", footer_xml)
        self.assertIn(" PAGE ", footer_xml)
        self.assertIn(" NUMPAGES ", footer_xml)
        self.assertEqual(footer_xml.count('w:fldCharType="begin"'), 2)
        self.assertEqual(footer_xml.count('w:fldCharType="separate"'), 2)
        self.assertEqual(footer_xml.count('w:fldCharType="end"'), 2)
        self.assertIn("2026-08-19", footer_xml)
        self.assertIn("校对稿", footer_xml)
        self.assertNotIn("<w:updateFields", settings_xml)

    def test_two_zone_header_uses_right_tab_and_bottom_border(self) -> None:
        document = parse_markdown(SAMPLE_MARKDOWN_WITH_TWO_ZONE_HEADER)
        payload = build_docx(
            document, load_preset("official-doc-cn-system-fonts-12pt")
        )

        with zipfile.ZipFile(io.BytesIO(payload)) as archive:
            header_xml = archive.read("word/header1.xml").decode("utf-8")
            footer_xml = archive.read("word/footer1.xml").decode("utf-8")

        self.assertIn("杭州红帆智药科技有限公司", header_xml)
        self.assertIn("内部资料，请勿外传", header_xml)
        self.assertEqual(header_xml.count("<w:tab/>"), 1)
        self.assertIn('w:val="right"', header_xml)
        self.assertNotIn('w:val="center"', header_xml)
        self.assertNotIn('w:val="clear"', header_xml)
        self.assertIn('w:sz w:val="21"', header_xml)
        self.assertIn('w:sz w:val="21"', footer_xml)
        self.assertIn('w:cs="Times New Roman"', header_xml)
        self.assertIn('w:cs="Times New Roman"', footer_xml)
        self.assertIn('w:szCs w:val="21"', header_xml)
        self.assertIn('w:szCs w:val="21"', footer_xml)
        self.assertRegex(
            footer_xml,
            r'<w:pPr>.*?<w:rPr><w:rFonts[^>]*w:ascii="Times New Roman"'
            r'[^>]*w:hAnsi="Times New Roman"[^>]*w:eastAsia="宋体"'
            r'[^>]*w:cs="Times New Roman"',
        )
        self.assertRegex(
            header_xml,
            r'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="000000"',
        )
        self.assertEqual(footer_xml.count("<w:tab/>"), 0)
        self.assertIn('<w:jc w:val="center"', footer_xml)
        self.assertIn(" PAGE ", footer_xml)
        self.assertEqual(footer_xml.count('w:fldCharType="begin"'), 1)
        self.assertEqual(footer_xml.count('w:fldCharType="separate"'), 1)
        self.assertEqual(footer_xml.count('w:fldCharType="end"'), 1)

        word_document = WordDocument(io.BytesIO(payload))
        header_style = word_document.styles["Header"]
        footer_style = word_document.styles["Footer"]
        header_character_style = word_document.styles["Header Char"]
        footer_character_style = word_document.styles["Footer Char"]
        self.assertEqual(header_style.font.name, "Times New Roman")
        self.assertEqual(footer_style.font.name, "Times New Roman")
        self.assertEqual(header_character_style.font.name, "Times New Roman")
        self.assertEqual(footer_character_style.font.name, "Times New Roman")
        self.assertEqual(header_style.font.size.pt, 10.5)
        self.assertEqual(footer_style.font.size.pt, 10.5)
        self.assertIsNone(header_style.element.pPr.find(qn("w:tabs")))
        self.assertIsNone(footer_style.element.pPr.find(qn("w:tabs")))
        self.assertEqual(
            header_style.element.rPr.rFonts.get(qn("w:eastAsia")),
            "仿宋",
        )
        self.assertEqual(
            footer_style.element.rPr.rFonts.get(qn("w:eastAsia")),
            "宋体",
        )
        self.assertEqual(
            footer_character_style.element.rPr.rFonts.get(qn("w:ascii")),
            "Times New Roman",
        )

    def test_page_content_rejects_style_fields_in_frontmatter(self) -> None:
        document = parse_markdown(
            "---\nwatermark:\n  text: 草案\n  opacity: 0.5\n---\n正文"
        )

        with self.assertRaisesRegex(ValueError, "Unsupported 'watermark' fields"):
            build_docx(document, load_preset("official-doc-cn"))

    def test_cli_can_print_preset_rules(self) -> None:
        stream = io.StringIO()
        with redirect_stdout(stream):
            exit_code = main(["--show-preset-rules", "official-doc-cn"])

        self.assertEqual(exit_code, 0)
        self.assertIn("结构化样式表", stream.getvalue())

    def test_cli_reports_version_schema_and_resolved_preset(self) -> None:
        version_output = io.StringIO()
        with self.assertRaises(SystemExit) as exit_context, redirect_stdout(version_output):
            main(["--version"])
        self.assertEqual(exit_context.exception.code, 0)
        self.assertEqual(version_output.getvalue().strip(), f"mdstyledocx {__version__}")

        schema_output = io.StringIO()
        with redirect_stdout(schema_output):
            self.assertEqual(main(["--show-preset-schema"]), 0)
        self.assertEqual(json.loads(schema_output.getvalue())["title"], "mdstyledocx preset")

        preset_output = io.StringIO()
        with redirect_stdout(preset_output):
            self.assertEqual(
                main(["--show-preset-json", "official-doc-cn-12pt"]), 0
            )
        definition = json.loads(preset_output.getvalue())
        self.assertEqual(definition["paragraph_defaults"]["line"], 400)
        self.assertNotIn("extends", definition)
